from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.core.exceptions import ValidationError
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import get_user_model
from .forms import RegistrationForm
from django.contrib.auth.decorators import user_passes_test

from django.contrib.auth.decorators import login_required
from .models import Praktiki, PeriodyPraktiki, Specialnosti, SpecialnostiGruppy, UchebnyjPlan, ProfessionalnyeModuli, Tip


def index(request):
    return render(request, 'main/index.html')

def register(request):
    return render(request, 'main/register.html')

def login_view(request):
    return render(request, 'main/login.html')

def admin_panel(request):
    return render(request, 'main/admin_panel.html')

def user_panel(request):
    return render(request, 'main/user_panel.html')

def search(request):
    query = request.GET.get('q', '')
    results = []  # Здесь можно добавить логику поиска (например, из базы данных)
    return render(request, 'main/search_results.html', {'query': query, 'results': results})

# Аутентификация
def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.username = form.cleaned_data['email']  # Используем email как логин
            user.set_password(form.cleaned_data['password1'])
            user.save()
            login(request, user)
            return redirect('login')  # Перенаправляем на страницу входа
    else:
        form = RegistrationForm()
    return render(request, 'main/register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']

        # Аутентификация пользователя
        user = authenticate(request, username=email, password=password)
        if user is not None:
            login(request, user)
            return redirect('index')  # Перенаправляем на главную страницу
        else:
            return render(request, 'main/login.html', {'error': 'Неверный логин или пароль'})
    return render(request, 'main/login.html')

def logout_view(request):
    logout(request)
    return redirect('index')

# админ
def admin_login(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']

        # Аутентификация пользователя
        user = authenticate(request, username=email, password=password)
        if user is not None and user.is_superuser:
            login(request, user)
            return redirect('/admin/')  # Перенаправляем на админ-панель
        else:
            return render(request, 'main/admin_login.html', {'error': 'Неверный логин или пароль администратора'})

    return render(request, 'main/admin_login.html')

# для вывода данных
from django.shortcuts import get_object_or_404
from .models import Praktiki, Zadaniya, UchebnyjPlan
from .models import Praktiki, PeriodyPraktiki, Specialnosti, SpecialnostiGruppy, UchebnyjPlan, ProfessionalnyeModuli, Tip, Zadaniya

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import (Praktiki, PeriodyPraktiki, Specialnosti, 
                    SpecialnostiGruppy, UchebnyjPlan, 
                    ProfessionalnyeModuli, Zadaniya, RukovoditeliPraktik)

@login_required
def user_panel(request):
    # Получаем текущего пользователя
    user = request.user

    # Получаем все практики, связанные с пользователем через related_name
    praktiki = Praktiki.objects.filter(rukovoditeli__rukovoditel_nomer=user)

    # Собираем данные для каждой практики
    data = []
    for praktika in praktiki:
        period = praktika.period_nomer
        tip = period.tip_kod
        specialnost = praktika.specialnost_kod
        gruppy = SpecialnostiGruppy.objects.filter(specialnosti_kod=specialnost)
        modul = praktika.modul_nomer  # Это объект модели ProfessionalnyeModuli
        mdks = ProfessionalnyeModuli.objects.filter(modul_nomer=modul)

        # Форматируем даты
        data_nachala = period.data_nachala.strftime('%d.%m.%Y')
        data_okonchaniya = period.data_okonchaniya.strftime('%d.%m.%Y')

        # Генерируем название практики
        try:
            # Извлекаем номер модуля            
            modul_nomer = modul.naimenovanie_modulya.split()[0].split('.')[1]
        except (AttributeError, IndexError):
            modul_nomer = "Н/Д"  # Если номер модуля не найден, используем значение по умолчанию

        # Получаем первую группу из QuerySet
        gruppa = gruppy.first()
        praktika_nazvanie = f"Группа_{gruppa.nazvanie_gruppy if gruppa else 'Н/Д'}_{period.god}_Задание_ПП{modul_nomer}"

        # Добавляем данные в список
        data.append({
            'nomer_praktiki': praktika.nomer_praktiki,
            'praktika_nazvanie': praktika_nazvanie,
            'god': period.god,
            'naimenovanie_tipa': tip.naimenovanie_tipa if tip else None,
            'kod_specialnosti': specialnost.kod_specialnosti,
            'nazvanie_specialnosti': specialnost.nazvanie_specialnosti,
            'nazvanie_gruppy': ', '.join([gruppa.nazvanie_gruppy for gruppa in gruppy]) if gruppy.exists() else 'Н/Д',
            'naimenovanie_modulya': modul.naimenovanie_modulya if modul and modul.naimenovanie_modulya else 'Н/Д',
            'naimenovanie_mdk': ', '.join([mdk.naimenovanie_mdk for mdk in mdks]) if mdks.exists() else 'Н/Д',
            'obem_v_chasah': period.obem_v_chasah,
            'data_nachala': data_nachala,
            'data_okonchaniya': data_okonchaniya,
            'praktika_id': praktika.nomer_praktiki,  # Для ссылки на детальную страницу
        })

    # Передаем данные в шаблон
    return render(request, 'main/user_panel.html', {'data': data})

def praktika_detail(request, pk):
    # Получаем практику по ее номеру
    praktika = get_object_or_404(Praktiki, nomer_praktiki=pk)
    # Проверяем, что modul_nomer существует
    if not praktika.modul_nomer:
        return render(request, 'main/praktika_detail.html', {
            'error': 'Модуль для данной практики не найден.'
        })
    # Получаем связанные данные
    period = praktika.period_nomer
    uchebnyj_plan = praktika.modul_nomer  # Это объект модели UchebnyjPlan
    mdks = ProfessionalnyeModuli.objects.filter(modul_nomer=uchebnyj_plan.nomer_modulya)
    zadaniya = Zadaniya.objects.filter(praktika_nomer=praktika)
    # Формируем название практики
    gruppa = SpecialnostiGruppy.objects.filter(specialnosti_kod=praktika.specialnost_kod).first()
    praktika_nazvanie = f"Группа_{gruppa.nazvanie_gruppy if gruppa else 'Н/Д'}_{period.god}_Задание_ПП{uchebnyj_plan.nomer_modulya.split('_')[-1]}"
    # Собираем данные для таблицы
    table_data = []
    for zadanie in zadaniya:
        kompetencii = uchebnyj_plan.kompetencii or 'Н/Д'
        table_data.append({
            'nazvanie_temy': zadanie.nazvanie_temy,
            'vidy_rabot': zadanie.vidy_rabot,
            'kompetencii': kompetencii,
            'kolichestvo_chasov_na_odnu_rabotu': zadanie.kolichestvo_chasov_na_odnu_rabotu,
        })
    return render(request, 'main/praktika_detail.html', {
        'praktika': praktika,
        'praktika_nazvanie': praktika_nazvanie,
        'table_data': table_data,
    })

from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Cm  # For font size and margins
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For paragraph alignment
from .models import Praktiki, PeriodyPraktiki, SpecialnostiGruppy, ProfessionalnyeModuli, Zadaniya, RukovoditeliPraktik
import datetime
from docx.enum.text import WD_COLOR_INDEX

def generate_praktika_doc(request, pk):
    # Retrieve practice data
    praktika = get_object_or_404(Praktiki, nomer_praktiki=pk)
    period = praktika.period_nomer
    specialnost = praktika.specialnost_kod
    gruppa = SpecialnostiGruppy.objects.filter(specialnosti_kod=specialnost).first()
    modul = praktika.modul_nomer
    mdks = ProfessionalnyeModuli.objects.filter(modul_nomer=modul)
    zadaniya = Zadaniya.objects.filter(praktika_nomer=praktika)

    try:
        # Extract module number
        modul_nomer = modul.naimenovanie_modulya.split()[0].split('.')[1]
    except (AttributeError, IndexError):
        modul_nomer = "Н/Д"  # If module number not found, use default value

    # Prepare data for the document
    data = {
        'praktika_nazvanie': f"Группа_{gruppa.nazvanie_gruppy if gruppa else 'Н/Д'}_{period.god}_Задание_ПП{modul_nomer}",
        'god': period.god,
        'naimenovanie_tipa': period.tip_kod.naimenovanie_tipa if period.tip_kod else 'Н/Д',
        'kod_specialnosti': specialnost.kod_specialnosti,
        'nazvanie_specialnosti': specialnost.nazvanie_specialnosti,
        'nazvanie_gruppy': gruppa.nazvanie_gruppy if gruppa else 'Н/Д',
        'naimenovanie_modulya': modul.naimenovanie_modulya if modul and modul.naimenovanie_modulya else 'Н/Д',
        'naimenovanie_mdk': ', '.join([mdk.naimenovanie_mdk for mdk in mdks]) if mdks.exists() else 'Н/Д',
        'obem_v_chasah': period.obem_v_chasah,
        'data_nachala': period.data_nachala.strftime('%d.%m.%Y'),
        'data_okonchaniya': period.data_okonchaniya.strftime('%d.%m.%Y'),
        'rukovoditeli': RukovoditeliPraktik.objects.filter(praktika_nomer=praktika),
        'zadaniya': zadaniya,
    }

    # Create a new Word document
    doc = Document()

    # Set default font to Times New Roman, size 11
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

    # Add heading "УТВЕРЖДАЮ"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("УТВЕРЖДАЮ")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add remaining text with specific formatting
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Заместитель директора по УПР")
    run.underline = True
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("ГБПОУ МО «Люберецкий техникум имени Героя Советского Союза, летчика-космонавта Ю.А. Гагарина»")
    run.underline = True
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"__________________ Т.И. Чиркунова")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"«___» __________________ {data['god']} г.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adjust indentation for the entire block
    for para in doc.paragraphs[-5:]:
        para.paragraph_format.left_indent = Cm(12)  # 12 cm indentation
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add individual assignment details
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Индивидуальное задание")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"на {'учебную' if data['naimenovanie_tipa'] == 'учебн' else 'производственную'} практическую подготовку")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"студенту ГБПОУ МО «Люберецкий техникум имени Героя Советского Союза, летчика-космонавта Ю.А. Гагарина»")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("___________________________________________________________________,")  # Placeholder for student's name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("ФИО")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"обучающемуся по специальности {data['kod_specialnosti']} \"{data['nazvanie_specialnosti']}\"")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Наименование профессионального модуля:")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(data['naimenovanie_modulya'])
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"включающему в себя:")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(data['naimenovanie_mdk'])
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Объем в часах: {data['obem_v_chasah']} часов")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Сроки прохождения практической подготовки: {data['data_nachala']} – {data['data_okonchaniya']}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"В ходе выполнения индивидуального задания в период {'учебной' if data['naimenovanie_tipa'] == 'учебн' else 'производственной'} практической подготовки студент должен освоить умения и приобрести общие и профессиональные компетенции.")

    # Add task table
    table_data = []
    total_hours = 0
    uchebnyj_plan = praktika.modul_nomer  # This is an object of model UchebnyjPlan
    for zadanie in data['zadaniya']:
        kompetencii = uchebnyj_plan.kompetencii or 'Н/Д'
        table_data.append({
            'nazvanie_temy': zadanie.nazvanie_temy,
            'vidy_rabot': zadanie.vidy_rabot,
            'kompetencii': kompetencii,
            'kolichestvo_chasov_na_odnu_rabotu': zadanie.kolichestvo_chasov_na_odnu_rabotu,
        })
        total_hours += zadanie.kolichestvo_chasov_na_odnu_rabotu

    # Add extra row for defense
    table_data.append({
        'nazvanie_temy': 'Защита отчетов по практике. Комплексный дифференцированный зачёт.',
        'vidy_rabot': '',
        'kompetencii': '',
        'kolichestvo_chasov_na_odnu_rabotu': 2,
    })
    total_hours += 2
    table_data.append({
        'nazvanie_temy': 'ВСЕГО',
        'vidy_rabot': '',
        'kompetencii': '',
        'kolichestvo_chasov_na_odnu_rabotu': total_hours,
    })

    # Add table to the document
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название темы'
    hdr_cells[1].text = 'Виды работ'
    hdr_cells[2].text = 'Профессиональные компетенции'
    hdr_cells[3].text = 'Кол-во часов'

    # Populate table rows
    # Сначала собираем все строки данных
    rows = []
    for row in table_data:
        rows.append({
            'nazvanie_temy': row['nazvanie_temy'],
            'vidy_rabot': row['vidy_rabot'],
            'kompetencii': row['kompetencii'],
            'kolichestvo_chasov_na_odnu_rabotu': str(row['kolichestvo_chasov_na_odnu_rabotu'])
        })

    # Добавляем строки в таблицу
    for row_data in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data['nazvanie_temy']
        row_cells[1].text = row_data['vidy_rabot']
        row_cells[3].text = row_data['kolichestvo_chasov_na_odnu_rabotu']

    # Теперь обрабатываем столбец с компетенциями
    if rows:
        # Получаем первую ячейку с компетенциями
        first_row = table.rows[1].cells[2]
        first_row.text = rows[0]['kompetencii']
        
        # Если есть более одной строки, объединяем все ячейки компетенций
        if len(rows) > 1:
            last_row = table.rows[len(rows)].cells[2]
            first_row.merge(last_row)

    

    # Add total hours
    # Add deadline for report submission
    paragraph = doc.add_paragraph("")
    paragraph = doc.add_paragraph(f"Срок предоставления Отчета к защите «{data['data_okonchaniya']}»")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add supervisors
    supervisors = data['rukovoditeli']
    supervisor_names = [
        f'{supervisor.rukovoditel_nomer.last_name} {supervisor.rukovoditel_nomer.first_name[0]}.{supervisor.rukovoditel_nomer.patronymic[0] if supervisor.rukovoditel_nomer.patronymic else ""}'
        for supervisor in supervisors
    ]
    doc.add_paragraph('Руководители практической подготовки')
    for name in supervisor_names:
        paragraph = doc.add_paragraph(f'____________________________ / {name} /')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph = doc.add_paragraph('ФИО руководителя')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add start date
    start_date = data['data_nachala'] if data['naimenovanie_tipa'] == 'учебн' else (
        datetime.datetime.strptime(data['data_nachala'], '%d.%m.%Y') - datetime.timedelta(days=1)
    ).strftime('%d.%m.%Y')
    paragraph = doc.add_paragraph(f'«{start_date}»')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{data["praktika_nazvanie"]}.docx"'
    doc.save(response)
    return response